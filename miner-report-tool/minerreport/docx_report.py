"""ساخت گزارش وورد/پی‌دی‌اف بر پایه‌ی قالب (سربرگ و کادرهای گزارش نمونه)."""

from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .pipeline import Device, ReportHeader, device_block


def _clear_run(run_element) -> None:
    for child in list(run_element):
        if child.tag != qn("w:rPr"):
            run_element.remove(child)


def _make_text_element(text: str):
    from docx.oxml import OxmlElement

    element = OxmlElement("w:t")
    element.set(qn("xml:space"), "preserve")
    element.text = text
    return element


def _accent_run(paragraph: Paragraph):
    """ران رنگی/زیرخط‌دار قالب (در نمونه: عدد مدت کارکرد)."""
    for run in paragraph.runs:
        color = run.font.color
        if run.underline or (color is not None and color.type is not None):
            return run._r
    return None


def _set_lines(paragraph: Paragraph, lines: list[str]) -> None:
    """چند خط را در یک پاراگراف با w:br می‌نویسد و قالب‌بندی خط اول را نگه می‌دارد."""
    from docx.oxml import OxmlElement

    if not paragraph.runs:
        paragraph.add_run("")
    mold = deepcopy(paragraph.runs[0]._r)
    _clear_run(mold)
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)
    for position, line in enumerate(lines):
        run_element = deepcopy(mold)
        run_element.append(_make_text_element(line))
        if position != len(lines) - 1:
            run_element.append(OxmlElement("w:br"))
        paragraph._p.append(run_element)


def _set_device_lines(paragraph: Paragraph, lines: list[str]) -> None:
    """مانند _set_lines، با این تفاوت که عدد مدت کارکرد قالب رنگی نمونه را می‌گیرد."""
    from docx.oxml import OxmlElement

    accent = _accent_run(paragraph)
    if accent is None or not paragraph.runs:
        _set_lines(paragraph, lines)
        return
    accent = deepcopy(accent)
    _clear_run(accent)
    default = deepcopy(paragraph.runs[0]._r)
    _clear_run(default)
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)

    for position, line in enumerate(lines):
        match = re.match(r"^(مدت کارکرد: )(\S+)( روز)$", line)
        segments = (
            [(match.group(1), False), (match.group(2), True), (match.group(3), False)]
            if match
            else [(line, False)]
        )
        for text, is_accent in segments:
            element = deepcopy(accent if is_accent else default)
            element.append(_make_text_element(text))
            paragraph._p.append(element)
        if position != len(lines) - 1:
            trailing = deepcopy(default)
            trailing.append(OxmlElement("w:br"))
            paragraph._p.append(trailing)


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    _set_lines(paragraph, [text])


def _cell_paragraph(cell) -> Paragraph:
    """پاراگراف اول سلول؛ پاراگراف‌های اضافی قالب حذف می‌شوند."""
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)
    return cell.paragraphs[0]


def _general_lines(header: ReportHeader) -> list[str]:
    return [
        "مشخصات کلی",
        f"پرونده - واحد : {header.case_unit}",
        f"آدرس : {header.address}",
        f"تعداد دستگاه بررسی شده: {header.device_count}",
        f"شماره گزارش: {header.report_number}",
    ]


def _report_info_lines(header: ReportHeader) -> list[str]:
    pad = " " * 72
    return [
        f"{pad}مشخصات گزارش",
        f"تاریخ گزارش: {header.report_date}",
        f"تائید مسئول مربوطه: {header.approver}",
    ]


def build_docx(
    header: ReportHeader,
    devices: list[Device],
    template: Path,
    output: Path,
) -> Path:
    document = Document(str(template))
    body = document.element.body
    children = list(body)

    title_paragraph = Paragraph(children[1], document)
    discovery_paragraph = Paragraph(children[3], document)
    notes_paragraph = Paragraph(children[8], document)
    signature_paragraph = deepcopy(children[6])

    general_table = document.tables[0]
    device_table = document.tables[1]
    device_row_mold = deepcopy(device_table.rows[0]._tr)
    report_table_element = deepcopy(device_table._tbl)
    report_row_mold = deepcopy(device_table.rows[3]._tr)

    _set_paragraph_text(
        title_paragraph,
        f" گزارش تخلیه اطلاعاتی دستگاه های ماینر شرکت توزیع نیروی برق استان {header.province}   ",
    )
    _set_paragraph_text(discovery_paragraph, f"تاریخ کشف: {header.discovery_date}")
    _set_lines(_cell_paragraph(general_table.cell(0, 0)), _general_lines(header))

    notes_text = "توضیحات کارشناسی: "
    if header.expert_notes:
        notes_text += header.expert_notes
    _set_paragraph_text(notes_paragraph, notes_text)

    # کادرهای دستگاه‌ها: یک ردیف برای هر دستگاه در جدول اول
    for row in list(device_table.rows):
        device_table._tbl.remove(row._tr)
    for device in devices:
        row_element = deepcopy(device_row_mold)
        device_table._tbl.append(row_element)
        row = device_table.rows[-1]
        _set_device_lines(
            _cell_paragraph(row.cells[0]),
            device_block(device, header.discovery_date).split("\n"),
        )

    # حذف کادرها و تیترهای اضافی قالب (دستگاه‌های ۳ تا ۱۱ و امضای شناور)
    for index in (11, 10, 9, 7, 6):
        body.remove(children[index])

    report_table = Table(report_table_element, document)
    for row in list(report_table.rows):
        report_table_element.remove(row._tr)
    report_table_element.append(deepcopy(report_row_mold))
    _set_lines(
        _cell_paragraph(report_table.rows[0].cells[0]),
        _report_info_lines(header),
    )

    body.append(report_table_element)
    body.append(signature_paragraph)  # امضای شناور، پایین کادر مشخصات گزارش
    body.append(children[13])  # sectPr باید آخرین عنصر body بماند

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    return output


def convert_to_pdf(docx_path: Path, output_dir: Path | None = None) -> Path:
    """تبدیل به PDF با LibreOffice (روی ویندوز هم soffice نصب باشد کار می‌کند)."""
    output_dir = output_dir or docx_path.parent
    binary = shutil.which("soffice") or shutil.which("libreoffice")
    if not binary:
        raise RuntimeError("LibreOffice نصب نیست؛ برای خروجی PDF آن را نصب کنید.")
    with tempfile.TemporaryDirectory() as profile:
        subprocess.run(
            [
                binary,
                f"-env:UserInstallation=file://{profile}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
        )
    pdf_path = output_dir / (docx_path.stem + ".pdf")
    if not pdf_path.exists():
        raise RuntimeError("تبدیل به PDF ناموفق بود.")
    return pdf_path
